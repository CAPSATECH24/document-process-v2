import time
import google.generativeai as genai
import streamlit as st
from tqdm import tqdm
import PyPDF2
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import RGBColor
import re
from docx.oxml import OxmlElement
from functools import wraps
import base64
from io import BytesIO
import requests
import threading
import queue
import os

# Configuración del registro de errores
logging.basicConfig(filename='error_log.txt', level=logging.ERROR,
                    format='%(asctime)s:%(levelname)s:%(message)s')

# Variables para el control de límites
requests_count = 0
tokens_used = 0
start_time = time.time()

# Variables globales para memoria de contexto
context_memory = {}

class ContextManager:
    def __init__(self):
        self.context = {}
        self.hierarchy = {}
        self.process_history = []
        
    def add_context(self, section_id, context_data):
        """Agrega o actualiza el contexto para una sección específica"""
        self.context[section_id] = {
            'data': context_data,
            'timestamp': time.time(),
            'dependencies': [],
            'references': []
        }
    
    def get_context(self, section_id):
        """Obtiene el contexto de una sección con sus dependencias"""
        if section_id in self.context:
            return self.context[section_id]
        return None
    
    def add_dependency(self, section_id, dependent_section):
        """Registra dependencias entre secciones"""
        if section_id in self.context:
            self.context[section_id]['dependencies'].append(dependent_section)
    
    def track_process(self, section_id, action):
        """Registra el historial de procesamiento"""
        self.process_history.append({
            'section': section_id,
            'action': action,
            'timestamp': time.time()
        })

class ProcessStructure:
    def __init__(self):
        self.sections = {}
        self.validation_rules = {}
        
    def add_section(self, section_id, parent_id=None):
        """Agrega una nueva sección a la estructura"""
        self.sections[section_id] = {
            'parent': parent_id,
            'children': [],
            'content': None,
            'status': 'pending'
        }
        if parent_id and parent_id in self.sections:
            self.sections[parent_id]['children'].append(section_id)
    
    def validate_section(self, section_id, content):
        """Valida el contenido de una sección según reglas predefinidas"""
        if section_id not in self.validation_rules:
            return True
        
        rules = self.validation_rules[section_id]
        return all(rule(content) for rule in rules)

def limit_control(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        model_version = kwargs.get('model_version', 'gemini-1.5-flash')
        check_limits(model_version)
        return func(*args, **kwargs)
    return wrapper

def check_limits(model_version='gemini-1.5-flash'):
    global requests_count, tokens_used, start_time
    current_time = time.time()

    # Definir límites según la versión del modelo
    limits = {
        'gemini-1.5-flash': (15, 1000000, 1500),
        'gemini-1.5-pro': (2, 32000, 50),
        'gemini-1.0-pro': (15, 32000, 1500)
    }

    RPM_LIMIT, TPM_LIMIT, RPD_LIMIT = limits.get(model_version, limits['gemini-1.5-flash'])

    # Reiniciar contador cada minuto
    if current_time - start_time > 60:
        requests_count = 0
        tokens_used = 0
        start_time = current_time

    # Comprobar límites de solicitudes por minuto
    if requests_count >= RPM_LIMIT:
        st.warning(f"Límite de {RPM_LIMIT} solicitudes por minuto alcanzado. Esperando 60 segundos...")
        time.sleep(60)
        requests_count = 0
        tokens_used = 0
        start_time = time.time()

    # Comprobar límites de tokens por minuto
    if tokens_used >= TPM_LIMIT:
        st.warning(f"Límite de {TPM_LIMIT} tokens por minuto alcanzado. Esperando 60 segundos...")
        time.sleep(60)
        requests_count = 0
        tokens_used = 0
        start_time = time.time()

def configurar_api(api_key):
    """
    Configura la clave API para la biblioteca google-generativeai.
    """
    if not api_key:
        raise ValueError("La clave API no puede estar vacía.")
    try:
        genai.configure(api_key=api_key)
        st.success("Configuración de la API exitosa.")
    except Exception as e:
        logging.error(f"Error al configurar la API: {e}")
        st.error(f"Error al configurar la API: {e}")
        raise

def crear_modelo(model_name='gemini-1.5-flash'):
    """
    Crea una instancia del modelo generativo especificado.
    """
    try:
        modelo = genai.GenerativeModel(model_name)
        st.success(f"Modelo '{model_name}' configurado correctamente.")
        return modelo
    except Exception as e:
        logging.error(f"Error al configurar el modelo generativo: {e}")
        st.error(f"Error al configurar el modelo generativo: {e}")
        raise

def cargar_archivo(uploaded_file):
    """
    Permite al usuario cargar un archivo .txt, .pdf o .docx y extrae su contenido.
    """
    if uploaded_file is not None:
        file_name = uploaded_file.name
        content = uploaded_file.read()
        try:
            text = process_file(file_name, content)
            st.success(f"Archivo '{file_name}' cargado exitosamente.")
            return text
        except Exception as e:
            logging.error(f"Error al leer el archivo: {e}")
            st.error(f"Error al leer el archivo: {e}")
            return None
    else:
        st.warning("No se cargó ningún archivo.")
        return None

def process_file(filename, content):
    """Procesa un solo archivo basado en su extensión."""
    if filename.lower().endswith('.txt'):
        text = content.decode('utf-8')
    elif filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(BytesIO(content))
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(BytesIO(content))
    else:
        raise ValueError(f"Formato de archivo no soportado: {filename}")
    return text

def extract_text_from_pdf(pdf_file):
    """
    Extrae texto de un archivo PDF.
    """
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        logging.error(f"Error al extraer texto del PDF: {e}")
        return ""

def extract_text_from_docx(docx_file):
    """
    Extrae texto de un archivo DOCX.
    """
    try:
        document = Document(docx_file)
        text = '\n'.join([para.text for para in document.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error al extraer texto del DOCX: {e}")
        return ""

def analizar_input_text(input_text):
    """
    Realiza un análisis semántico del texto de entrada para extraer entidades clave y contexto.
    """
    # Implementar lógica de análisis semántico si es necesario
    # Por simplicidad, aquí simplemente retornamos el texto tal cual
    return input_text

def validar_contenido(contenido, subpunto):
    """
    Valida el contenido generado con reglas más flexibles según el tipo de sección.
    """
    if not contenido or len(contenido.strip()) < 10:  # Validación mínima de contenido
        logging.warning(f"Contenido vacío o demasiado corto para {subpunto}")
        return False

    contenido_lower = contenido.lower()
    subpunto_lower = subpunto.lower()

    # Caso especial para la portada
    if "portada" in subpunto_lower:
        return len(contenido.strip()) >= 20  # Solo verificar que tenga un contenido mínimo
    
    # Caso especial para objetivos
    if "objetivo" in subpunto_lower:
        # Para objetivos, solo verificamos que el contenido sea coherente y no esté vacío
        patrones_prohibidos = ['no está claro', 'indefinido', 'etc.', 'error']
        return not any(patron in contenido_lower for patron in patrones_prohibidos)
    
    # Caso especial para secciones introductorias o de contexto
    if any(keyword in subpunto_lower for keyword in ['introducción', 'contexto', 'alcance']):
        # Validación más flexible para secciones descriptivas
        return len(contenido.strip()) >= 50  # Solo verificar longitud mínima razonable
    
    # Para el resto de las secciones, mantener algunas validaciones básicas
    criterios_basicos = {
        'longitud_minima': 30,  # Reducida significativamente
        'patrones_prohibidos': ['error', 'undefined', 'null']  # Solo patrones críticos
    }
    
    # Validar longitud mínima básica
    if len(contenido) < criterios_basicos['longitud_minima']:
        logging.warning(f"Contenido demasiado corto para {subpunto}")
        return False
    
    # Validar solo patrones críticos prohibidos
    if any(patron in contenido_lower for patron in criterios_basicos['patrones_prohibidos']):
        logging.warning(f"Patrones críticos prohibidos encontrados en {subpunto}")
        return False
    
    return True

def generar_subpunto(modelo, seccion, subpunto, instrucciones, input_text, context_manager, tipo_proceso, max_retries=5, model_version='gemini-1.5-flash'):
    """
    Genera el contenido para un subpunto específico con manejo más flexible del contenido,
    incluyendo notas explicativas importantes cuando sea necesario.
    """
    retry_count = 0
    while retry_count < max_retries:
        try:
            # Obtener contexto relevante
            contexto_actual = context_manager.get_context(f"{seccion}_{subpunto}")
            contexto_padre = context_manager.get_context(seccion)
            
            # Determinar el tipo de sección
            subpunto_lower = subpunto.lower()
            
            if "objetivo" in subpunto_lower:
                prompt = f"""
                Genera un objetivo claro y conciso para la sección:
                
                Sección: {seccion}
                Subpunto: {subpunto}
                Tipo de proceso: {tipo_proceso}
                
                El objetivo debe:
                1. Ser claro y directo
                2. Explicar el propósito principal
                3. Ser relevante para el proceso
                
                Incluye también:
                - Notas importantes sobre consideraciones especiales
                - Cualquier prerrequisito o condición relevante
                - Impacto esperado del proceso
                
                Contexto adicional: {input_text}
                
                Formato esperado:
                OBJETIVO: [texto del objetivo]
                
                NOTAS IMPORTANTES:
                - [Nota 1]
                - [Nota 2]
                ...
                """
            elif "portada" in subpunto_lower:
                prompt = f"""
                Genera el contenido para la portada:
                
                Incluye:
                1. Título descriptivo
                2. Fecha
                3. Versión o referencia
                
                Tipo de proceso: {tipo_proceso}
                """
            else:
                prompt = f"""
                Genera contenido detallado para la sección, incluyendo notas explicativas importantes:
                
                Sección: {seccion}
                Subpunto: {subpunto}
                Tipo: {tipo_proceso}
                
                Instrucciones: {instrucciones}
                
                El contenido debe incluir:
                1. Descripción clara del proceso o subproceso
                2. Pasos detallados cuando sea aplicable
                3. NOTAS IMPORTANTES que incluyan:
                   - Consideraciones especiales
                   - Advertencias o precauciones
                   - Mejores prácticas
                   - Casos excepcionales
                   - Dependencias con otros procesos
                4. Referencias a documentación relacionada si existe
                
                Contexto previo: {contexto_padre['data'] if contexto_padre else 'No disponible'}
                
                Información adicional: {input_text}
                
                Formato esperado:
                [CONTENIDO PRINCIPAL]
                
                NOTAS IMPORTANTES:
                - [Nota 1]
                - [Nota 2]
                ...
                
                REFERENCIAS:
                - [Referencia 1]
                - [Referencia 2]
                ...
                """
            
            # Generar respuesta
            response = modelo.generate_content(prompt)
            contenido = response.text
            
            # Validar y formatear el contenido
            if validar_contenido(contenido, subpunto):
                # Procesar y formatear las notas si existen
                contenido_formateado = formatear_contenido_con_notas(contenido)
                context_manager.add_context(f"{seccion}_{subpunto}", contenido_formateado)
                context_manager.track_process(f"{seccion}_{subpunto}", "generacion_exitosa")
                return contenido_formateado
            
            retry_count += 1
            if retry_count < max_retries:
                time.sleep(1)
            
        except Exception as e:
            logging.error(f"Error en generación de subpunto {subpunto}: {str(e)}")
            retry_count += 1
            if retry_count < max_retries:
                time.sleep(1)
    
    if 'contenido' in locals():
        logging.warning(f"Retornando último contenido generado para {subpunto} a pesar de no pasar validación")
        return contenido
    
    raise Exception(f"No se pudo generar contenido válido para el subpunto {subpunto} después de {max_retries} intentos")

def formatear_contenido_con_notas(contenido):
    """
    Formatea el contenido asegurando que las notas importantes estén correctamente resaltadas.
    """
    secciones = contenido.split('\n\n')
    contenido_formateado = []
    
    for seccion in secciones:
        if 'NOTAS IMPORTANTES:' in seccion:
            # Formatear las notas con viñetas y resaltado
            notas = seccion.split('\n')
            notas_formateadas = [notas[0]]  # Mantener el título
            for nota in notas[1:]:
                if nota.strip():
                    if not nota.strip().startswith('-'):
                        nota = f"- {nota.strip()}"
                    notas_formateadas.append(f"{nota}")
            contenido_formateado.append('\n'.join(notas_formateadas))
        else:
            contenido_formateado.append(seccion)
    
    return '\n\n'.join(contenido_formateado)

def eliminar_redundancias(texto):
    """
    Elimina oraciones o párrafos redundantes del texto proporcionado.
    """
    # Implementación simple para eliminar líneas duplicadas
    lines = texto.split('\n')
    seen = set()
    new_lines = []
    for line in lines:
        if line.strip() not in seen:
            seen.add(line.strip())
            new_lines.append(line)
    return '\n'.join(new_lines)

def parsear_contenido_para_word(contenido):
    """
    Parsea el contenido generado para aplicar formato en Word.
    """
    lines = contenido.split('\n')
    return lines

def process_line_with_formatting(paragraph, text):
    """
    Procesa una línea de texto, aplicando formato de negrita y cursiva donde se indique.
    """
    # Expresiones regulares para negrita y cursiva
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'\*(.*?)\*')

    # Reemplazar negritas y cursivas
    def replace_bold(match):
        run = paragraph.add_run(match.group(1))
        run.bold = True
        return ''

    def replace_italic(match):
        run = paragraph.add_run(match.group(1))
        run.italic = True
        return ''

    # Procesar el texto
    pos = 0
    while pos < len(text):
        bold_match = bold_pattern.search(text, pos)
        italic_match = italic_pattern.search(text, pos)

        next_match = None
        is_bold = False
        if bold_match and (not italic_match or bold_match.start() <= italic_match.start()):
            next_match = bold_match
            is_bold = True
        elif italic_match:
            next_match = italic_match
            is_bold = False

        if next_match:
            # Añadir texto antes del match
            if next_match.start() > pos:
                paragraph.add_run(text[pos:next_match.start()])
            # Añadir texto formateado
            if is_bold:
                replace_bold(next_match)
            else:
                replace_italic(next_match)
            pos = next_match.end()
        else:
            # Añadir el resto del texto
            paragraph.add_run(text[pos:])
            break

def agregar_contenido_al_documento(documento, contenido, tipo_proceso=None):
    """
    Agrega el contenido formateado al documento Word, incluyendo ejemplos si es necesario.
    """
    lines = parsear_contenido_para_word(contenido)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith('# '):
            titulo = line.replace('# ', '').strip()
            encabezado = documento.add_heading(titulo, level=1)
            encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            titulo = line.replace('## ', '').strip()
            encabezado = documento.add_heading(titulo, level=2)
            encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('### '):
            titulo = line.replace('### ', '').strip()
            documento.add_heading(titulo, level=3)
        elif line.startswith('#### '):
            titulo = line.replace('#### ', '').strip()
            documento.add_heading(titulo, level=4)
        elif line.startswith('- '):
            # Bullet list
            p = documento.add_paragraph(style='List Bullet')
            process_line_with_formatting(p, line[2:])
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif re.match(r'\s{2,}- ', line):
            # Sub-bullet list
            p = documento.add_paragraph(style='List Bullet 2')
            process_line_with_formatting(p, line.strip()[2:])
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif re.match(r'\d+\.', line):
            # Numbered list
            p = documento.add_paragraph(style='List Number')
            process_line_with_formatting(p, line[line.find('.')+1:].strip())
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif re.match(r'\s{2,}\d+\.', line):
            # Sub-numbered list
            p = documento.add_paragraph(style='List Number 2')
            process_line_with_formatting(p, line.strip()[line.strip().find('.')+1:].strip())
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif line.startswith('|'):
            # Tabla: Extraer todas las líneas de la tabla
            tabla = []
            while i < len(lines) and lines[i].startswith('|'):
                tabla.append(lines[i])
                i += 1
            crear_tabla(documento, tabla)
            continue  # Ya hemos incrementado 'i'
        elif line.startswith('```mermaid'):
            # Diagrama Mermaid
            mermaid_code = ''
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                mermaid_code += lines[i] + '\n'
                i += 1
            generar_diagrama_mermaid(documento, mermaid_code)
        elif 'imagen:' in line.lower():
            # Insertar imagen
            imagen = line.split('imagen:')[1].strip()
            try:
                documento.add_picture(imagen, width=Inches(6))
                last_paragraph = documento.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                logging.error(f"Error al agregar la imagen {imagen}: {e}")
                p = documento.add_paragraph(f"No se pudo agregar la imagen {imagen}.")
        elif tipo_proceso == 'software' and 'Ejemplo' in line:
            p = documento.add_paragraph()
            p.add_run("Ejemplo:").bold = True
            p.add_run(" Para iniciar el proceso, abre el programa XYZ, haz clic en 'Archivo', selecciona 'Nuevo', y sigue las indicaciones en pantalla.")
        else:
            # Párrafo normal
            p = documento.add_paragraph()
            process_line_with_formatting(p, line)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        i += 1

def crear_tabla(documento, tabla):
    """
    Crea una tabla en el documento Word a partir de las líneas proporcionadas.
    """
    if len(tabla) < 2:
        return  # No hay suficientes líneas para una tabla

    # Procesar las líneas de la tabla
    headers = [h.strip() for h in tabla[0].strip('|').split('|')]
    data_rows = [row.strip('|').split('|') for row in tabla[1:] if row.strip()]

    # Crear la tabla en Word
    table = documento.add_table(rows=1, cols=len(headers))
    table.style = 'EstiloTablaPersonalizado'
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Formatear la fila de encabezado
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header.strip()
        # Formatear el texto del encabezado
        paragraph = hdr_cells[idx].paragraphs[0]
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        # Establecer fondo gris claro para el encabezado
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Color gris claro
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Añadir filas de datos
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for idx, cell_text in enumerate(row_data):
            if idx < len(row_cells):
                cell = row_cells[idx]
                cell.text = cell_text.strip()
                # Alinear el texto
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = paragraph.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Aplicar sombreado a filas alternas
    for idx, row in enumerate(table.rows[1:], start=1):
        if idx % 2 == 0:
            for cell in row.cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F2F2')  # Color gris claro
                cell._tc.get_or_add_tcPr().append(shading_elm)

def simplificar_diagrama(mermaid_code):
    """
    Simplifica el diagrama Mermaid eliminando elementos redundantes.
    """
    # Implementar lógica de simplificación si es necesario
    # Por ejemplo, eliminar nodos duplicados o conexiones innecesarias
    return mermaid_code

def generar_diagrama_mermaid(documento, mermaid_code):
    """
    Genera una imagen de diagrama Mermaid y la inserta en el documento.
    """
    # Simplificar el diagrama
    mermaid_code = simplificar_diagrama(mermaid_code)

    # Codificar el código Mermaid en base64
    encoded_code = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
    # URL para generar el diagrama usando un servicio web (ejemplo: mermaid.ink)
    url = f"https://mermaid.ink/img/{encoded_code}"

    try:
        response = requests.get(url)
        if response.status_code == 200:
            image_data = BytesIO(response.content)
            documento.add_picture(image_data, width=Inches(6))
            last_paragraph = documento.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p = documento.add_paragraph('No se pudo generar el diagrama Mermaid.')
    except Exception as e:
        logging.error(f"Error al generar el diagrama Mermaid: {e}")
        p = documento.add_paragraph('Error al generar el diagrama Mermaid.')

def generar_indice_temporal(input_text):
    """
    Implementa una función que escanee y genere un índice temporal de todos los procesos principales y sus respectivos subprocesos.
    """
    # Utilizar expresiones regulares para extraer títulos numerados
    pattern = re.compile(r'^(\d+(?:\.\d+)*)(?:\s+|\.)\s*(.*)', re.MULTILINE)
    matches = pattern.findall(input_text)

    indice = ''
    for number, title in matches:
        indice += f"{number} {title}\n"

    return indice

def obtener_subpuntos(indice_seccion):
    """
    Extrae todos los subpuntos del índice de la sección y construye una estructura jerárquica.
    """
    lines = indice_seccion.strip().split('\n')
    subpuntos = []
    stack = []

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        # Match numbering (e.g., 1., 1.1., 1.1.1.)
        match = re.match(r'^(\d+(?:\.\d+)*)(?:\s+|\.)\s*(.*)', stripped_line)
        if match:
            number = match.group(1)
            title = match.group(2)
            level = number.count('.')
            subpunto = {
                'titulo': f"{number} {title}",
                'nivel': level,
                'subpuntos': []
            }

            while stack and stack[-1]['nivel'] >= level:
                stack.pop()
            if stack:
                stack[-1]['subpuntos'].append(subpunto)
            else:
                subpuntos.append(subpunto)
            stack.append(subpunto)
        else:
            # Si la línea no coincide con la numeración, se ignora
            continue

    return subpuntos

def generar_contenido_subpuntos(modelo, documento, seccion, subpuntos, instrucciones, input_text, context_manager, model_version, nivel):
    """
    Genera contenido recursivamente para cada subpunto y lo agrega al documento.
    """
    for subpunto in subpuntos:
        titulo = subpunto['titulo']
        encabezado = documento.add_heading(titulo, level=nivel)
        # Centrar encabezados principales
        if nivel <= 2:
            encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Determinar el tipo de proceso basado en palabras clave
        if any(keyword in titulo.lower() for keyword in ['software', 'sistema', 'aplicación']):
            tipo_proceso = 'software'
        elif any(keyword in titulo.lower() for keyword in ['administrativo', 'gestión', 'proceso']):
            tipo_proceso = 'administrativo'
        else:
            tipo_proceso = None

        # **Ajustar las instrucciones para la sección "6. Descripción Detallada de Procesos"**
        if '6.' in seccion or 'descripción detallada de procesos' in seccion.lower():
            instrucciones_subpunto = "Desarrolle un procedimiento detallado para los procesos, incluyendo instrucciones paso a paso, numeradas y con estructura, especialmente si se documentan procesos de programas o tecnología que lo necesiten."
        else:
            instrucciones_subpunto = instrucciones

        # Generar contenido para el subpunto
        texto_subpunto = generar_subpunto(modelo, seccion, titulo, instrucciones_subpunto, input_text, context_manager, tipo_proceso, model_version=model_version)
        if texto_subpunto.startswith("Error al generar"):
            p = documento.add_paragraph(texto_subpunto)
            p.style = 'Intense Quote'
        else:
            agregar_contenido_al_documento(documento, texto_subpunto, tipo_proceso)

        # Procesar subpuntos anidados
        if subpunto['subpuntos']:
            generar_contenido_subpuntos(modelo, documento, titulo, subpunto['subpuntos'], instrucciones_subpunto, input_text, context_manager, model_version, nivel + 1)

def add_table_of_contents(documento):
    """
    Agrega una Tabla de Contenidos al documento con formato mejorado.
    """
    # Agregar título para la tabla de contenidos
    paragraph = documento.add_paragraph()
    run = paragraph.add_run("Tabla de Contenidos")
    run.bold = True
    run.font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Agregar un párrafo para la tabla de contenidos
    paragraph = documento.add_paragraph()
    run = paragraph.add_run()
    
    # Crear la tabla de contenidos
    toc = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartObj.append(docPartGallery)
    sdtPr.append(docPartObj)
    toc.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar1)
    p.append(r)
    
    r = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Modificado para mantener niveles cerrados
    r.append(instrText)
    p.append(r)
    
    r = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r.append(fldChar2)
    p.append(r)
    
    sdtContent.append(p)
    toc.append(sdtContent)
    paragraph._p.append(toc)
    
    # Agregar un salto de página después de la tabla de contenidos
    documento.add_paragraph().add_run().add_break()

def configurar_estilos(documento):
    """
    Configura estilos personalizados para el documento.
    """
    styles = documento.styles

    # Establecer el idioma del documento a español
    documento.core_properties.language = 'es-ES'

    # Configurar estilo para el texto normal
    estilo_normal = styles['Normal']
    font = estilo_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    estilo_normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    estilo_normal.paragraph_format.line_spacing = 1.15
    estilo_normal.paragraph_format.space_before = Pt(6)
    estilo_normal.paragraph_format.space_after = Pt(6)

    # Configurar estilos para encabezados
    for i in range(1, 5):
        estilo = styles[f'Heading {i}']
        font = estilo.font
        font.name = 'Arial'
        font.size = Pt(16 - (i * 2))
        font.bold = True
        # Ajustar color
        if i == 1:
            font.color.rgb = RGBColor(0, 51, 102)
        elif i == 2:
            font.color.rgb = RGBColor(31, 73, 125)
        elif i == 3:
            font.color.rgb = RGBColor(79, 129, 189)
        # Ajustar formato de párrafo
        paragraph_format = estilo.paragraph_format
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)
        paragraph_format.keep_with_next = True
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i <= 2 else WD_PARAGRAPH_ALIGNMENT.LEFT
        estilo.paragraph_format.outline_level = i - 1  # Ajustar el nivel de esquema

    # Configurar estilos para listas
    estilo_lista_viñetas = styles['List Bullet']
    estilo_lista_viñetas.font.name = 'Arial'
    estilo_lista_viñetas.font.size = Pt(11)
    estilo_lista_viñetas.paragraph_format.left_indent = Inches(0.25)
    estilo_lista_viñetas.paragraph_format.space_before = Pt(0)
    estilo_lista_viñetas.paragraph_format.space_after = Pt(0)
    estilo_lista_viñetas.paragraph_format.line_spacing = 1.15
    estilo_lista_viñetas.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Verificar si 'List Bullet 2' ya existe antes de agregarlo
    if 'List Bullet 2' not in styles:
        estilo_lista_viñetas2 = styles.add_style('List Bullet 2', WD_STYLE_TYPE.PARAGRAPH)
        estilo_lista_viñetas2.base_style = estilo_lista_viñetas
        estilo_lista_viñetas2.paragraph_format.left_indent = Inches(0.5)
    else:
        estilo_lista_viñetas2 = styles['List Bullet 2']
        estilo_lista_viñetas2.paragraph_format.left_indent = Inches(0.5)

    # Configurar estilo para sublistas nivel 3
    if 'List Bullet 3' not in styles:
        estilo_lista_viñetas3 = styles.add_style('List Bullet 3', WD_STYLE_TYPE.PARAGRAPH)
        estilo_lista_viñetas3.base_style = estilo_lista_viñetas2
        estilo_lista_viñetas3.paragraph_format.left_indent = Inches(0.75)
    else:
        estilo_lista_viñetas3 = styles['List Bullet 3']
        estilo_lista_viñetas3.paragraph_format.left_indent = Inches(0.75)

    estilo_lista_numerada = styles['List Number']
    estilo_lista_numerada.font.name = 'Arial'
    estilo_lista_numerada.font.size = Pt(11)
    estilo_lista_numerada.paragraph_format.left_indent = Inches(0.25)
    estilo_lista_numerada.paragraph_format.space_before = Pt(0)
    estilo_lista_numerada.paragraph_format.space_after = Pt(0)
    estilo_lista_numerada.paragraph_format.line_spacing = 1.15
    estilo_lista_numerada.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Verificar si 'List Number 2' ya existe antes de agregarlo
    if 'List Number 2' not in styles:
        estilo_lista_numerada2 = styles.add_style('List Number 2', WD_STYLE_TYPE.PARAGRAPH)
        estilo_lista_numerada2.base_style = estilo_lista_numerada
        estilo_lista_numerada2.paragraph_format.left_indent = Inches(0.5)
    else:
        estilo_lista_numerada2 = styles['List Number 2']
        estilo_lista_numerada2.paragraph_format.left_indent = Inches(0.5)

    # Configurar estilo para sublistas numeradas nivel 3
    if 'List Number 3' not in styles:
        estilo_lista_numerada3 = styles.add_style('List Number 3', WD_STYLE_TYPE.PARAGRAPH)
        estilo_lista_numerada3.base_style = estilo_lista_numerada2
        estilo_lista_numerada3.paragraph_format.left_indent = Inches(0.75)
    else:
        estilo_lista_numerada3 = styles['List Number 3']
        estilo_lista_numerada3.paragraph_format.left_indent = Inches(0.75)

    # Estilo para el índice de contenidos
    if 'IndiceContenido' not in styles:
        estilo_indice = styles.add_style('IndiceContenido', WD_STYLE_TYPE.PARAGRAPH)
        estilo_indice.font.name = 'Arial'
        estilo_indice.font.size = Pt(11)
        estilo_indice.paragraph_format.space_before = Pt(0)
        estilo_indice.paragraph_format.space_after = Pt(0)
        estilo_indice.paragraph_format.line_spacing = 1.15
        estilo_indice.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:
        estilo_indice = styles['IndiceContenido']

    # Estilo personalizado para tablas
    if 'EstiloTablaPersonalizado' not in styles:
        estilo_tabla = styles.add_style('EstiloTablaPersonalizado', WD_STYLE_TYPE.TABLE)
        estilo_tabla.font.name = 'Arial'
        estilo_tabla.font.size = Pt(11)
    else:
        estilo_tabla = styles['EstiloTablaPersonalizado']

    # Configurar márgenes de página estándar
    sections = documento.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Añadir encabezados y pies de página con numeración de páginas
    for section in documento.sections:
        # Encabezado
        header = section.header
        header.is_linked_to_previous = False
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = "Documentación de Procesos ISO 9001"
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_paragraph.style = styles['Header']

        # Pie de página
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_paragraph.text = "Página "
        page_field = footer_paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        page_field._r.append(fldChar_begin)
        page_field._r.append(instrText)
        page_field._r.append(fldChar_end)

def agregar_recursos_adicionales(documento, enlaces):
    """
    Agrega una sección de recursos adicionales con enlaces útiles.
    """
    documento.add_heading('Recursos Adicionales', level=2)
    for titulo, url in enlaces.items():
        p = documento.add_paragraph()
        run = p.add_run(titulo)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True
        p.add_run(f": {url}")

def generar_documento(modelo, input_text, indice, model_version='gemini-1.5-flash'):
    """
    Genera el documento completo utilizando el índice proporcionado y el contenido de cada subpunto.
    """
    documento = Document()
    configurar_estilos(documento)

    # Añadir metadatos al documento
    documento.core_properties.author = 'Ingeniero Industrial Senior'
    documento.core_properties.title = 'Documentación de Procesos ISO 9001'
    documento.core_properties.keywords = 'ISO 9001, Procesos, Calidad'

    # Portada del Documento
    documento.add_section()
    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Añadir el logotipo de la empresa (si existe)
    try:
        documento.add_picture('logo_empresa.png', width=Inches(2))
        last_paragraph = documento.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        logging.warning(f"No se pudo agregar el logotipo: {e}")

    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Documentación de Procesos ISO 9001: Guía Integral de Implementación y Control de Calidad')
    run.bold = True
    font = run.font
    font.name = 'Arial'
    font.size = Pt(24)

    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Estandarización de Procesos y Control de Calidad para Nuevas Empresas')
    font = run.font
    font.name = 'Arial'
    font.size = Pt(18)

    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Autor: Ingeniero Industrial Senior (Nombre Completo y Credenciales)')
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Fecha de Elaboración: [Fecha]')
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Versión del Documento: 1.0')
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Salto de página después de la portada
    documento.add_page_break()

    # Agregar la Tabla de Contenidos
    documento.add_heading('Índice de Contenidos', level=1)
    add_table_of_contents(documento)

    # Salto de página después del índice
    documento.add_page_break()

    # Analizar el texto de entrada y extraer entidades clave
    input_analysis = analizar_input_text(input_text)

    # Inicializar memoria de contexto
    context_manager = ContextManager()

    # Generar el índice temporal a partir del input_text
    indice_temporal = generar_indice_temporal(input_text)
    if indice_temporal.strip() == '':
        indice_temporal = indice  # Usar el índice proporcionado si no se puede generar uno temporal

    # Parsear el índice y generar contenido
    subpuntos = obtener_subpuntos(indice_temporal)
    instrucciones_generales = "Desarrolle el contenido de la sección basándose en el texto de referencia y alineado con las normativas ISO 9001."

    for seccion in subpuntos:
        titulo_seccion = seccion['titulo']
        nivel_seccion = seccion['nivel'] + 1  # Ajustar el nivel inicial

        encabezado = documento.add_heading(titulo_seccion, level=nivel_seccion)
        if nivel_seccion <= 2:
            encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Determinar el tipo de proceso
        if any(keyword in titulo_seccion.lower() for keyword in ['software', 'sistema', 'aplicación']):
            tipo_proceso = 'software'
        elif any(keyword in titulo_seccion.lower() for keyword in ['administrativo', 'gestión', 'proceso']):
            tipo_proceso = 'administrativo'
        else:
            tipo_proceso = None

        # **Ajustar las instrucciones para la sección "6. Descripción Detallada de Procesos"**
        if '6.' in titulo_seccion or 'descripción detallada de procesos' in titulo_seccion.lower():
            instrucciones_seccion = "Desarrolle un procedimiento detallado para los procesos, incluyendo instrucciones paso a paso, numeradas y con estructura, especialmente si se documentan procesos de programas o tecnología que lo necesiten."
        else:
            instrucciones_seccion = instrucciones_generales

        # Generar contenido para la sección principal
        texto_seccion = generar_subpunto(modelo, titulo_seccion, titulo_seccion, instrucciones_seccion, input_analysis, context_manager, tipo_proceso, model_version=model_version)
        if texto_seccion.startswith("Error al generar"):
            p = documento.add_paragraph(texto_seccion)
            p.style = 'Intense Quote'
        else:
            agregar_contenido_al_documento(documento, texto_seccion, tipo_proceso)

        # Generar contenido para los subpuntos
        if seccion['subpuntos']:
            generar_contenido_subpuntos(modelo, documento, titulo_seccion, seccion['subpuntos'], instrucciones_seccion, input_analysis, context_manager, model_version, nivel_seccion + 1)

        # Agregar salto de página después de cada sección importante
        documento.add_page_break()

    # Agregar recursos adicionales (ejemplo)
    enlaces_recursos = {
        'Normativa ISO 9001': 'https://www.iso.org/iso-9001-quality-management.html',
        'Manual de Usuario del Software XYZ': 'https://www.softwarexyz.com/manual'
    }
    agregar_recursos_adicionales(documento, enlaces_recursos)

    return documento

def guardar_documento_word(documento, nombre_archivo):
    """
    Guarda el documento generado en un archivo Word.
    """
    try:
        documento.save(nombre_archivo)
        st.success(f"El documento ha sido guardado como '{nombre_archivo}'.")
        with open(nombre_archivo, "rb") as file:
            btn = st.download_button(
                label="Descargar documento",
                data=file,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        logging.error(f"Error al guardar el documento: {e}")
        st.error(f"Error al guardar el documento: {e}")

def definir_plantilla_contenido():
    """
    Define la estructura del documento con secciones y subsecciones.
    Retorna una lista de diccionarios que representan el esquema del documento.
    """
    plantilla = [
        {
            "seccion": "Portada",
            "subsecciones": ["Título", "Información General"]
        },
        {
            "seccion": "Información del Proceso",
            "subsecciones": [
                "Objetivo del Proceso",
                "Alcance",
                "Roles y Responsabilidades",
                "Definiciones Importantes"
            ]
        },
        {
            "seccion": "Descripción del Proceso",
            "subsecciones": [
                "Diagrama del Proceso",
                "Actividades Principales",
                "Subprocesos",
                "Puntos de Control"
            ]
        },
        {
            "seccion": "Elementos del Proceso",
            "subsecciones": [
                "Entradas",
                "Salidas",
                "Recursos Necesarios",
                "Sistemas Involucrados"
            ]
        },
        {
            "seccion": "Consideraciones Especiales",
            "subsecciones": [
                "Políticas Aplicables",
                "Restricciones",
                "Casos Excepcionales",
                "Mejores Prácticas"
            ]
        },
        {
            "seccion": "Métricas y Control",
            "subsecciones": [
                "Indicadores de Desempeño",
                "Puntos de Medición",
                "Acciones Correctivas"
            ]
        },
        {
            "seccion": "Documentación Relacionada",
            "subsecciones": [
                "Referencias",
                "Formatos Aplicables",
                "Documentos de Soporte"
            ]
        }
    ]
    return plantilla

def evaluar_calidad_contenido(contenido):
    """
    Evalúa la calidad del contenido basado en criterios predefinidos.
    Retorna un puntaje de calidad.
    """
    # Criterios de ejemplo: longitud mínima, presencia de palabras clave
    puntaje = 0
    if len(contenido) > 50:
        puntaje += 1
    if "clave" in contenido:
        puntaje += 1
    return puntaje

def generar_contenido_para_subseccion(seccion, subseccion, context_memory):
    """
    Genera contenido para una subsección específica utilizando información de contexto.
    Retorna un texto que representa el contenido generado.
    """
    # Obtener contexto relevante
    contexto = context_memory.get(seccion, {}).get(subseccion, "")
    contenido = f"Contenido generado para {subseccion} en la sección {seccion}."

    # Refinar contenido iterativamente
    for _ in range(3):  # Realizar tres iteraciones de refinamiento
        # Evaluar calidad del contenido
        calidad = evaluar_calidad_contenido(contenido)
        if calidad < 2:
            contenido += " (refinado para mejorar calidad)"

    return contenido

def main():
    """
    Función principal que coordina la generación del documento.
    """
    st.set_page_config(page_title="Generador de Documentación de Procesos", layout="wide")
    st.title("Generador de Documentación de Procesos")
    st.write("Esta aplicación genera automáticamente un documento de procesos alineado con la norma ISO 9001.")

    # Inicializar estados en la sesión si no existen
    if 'current_stage' not in st.session_state:
        st.session_state.current_stage = "Esperando configuración"
    if 'progress_value' not in st.session_state:
        st.session_state.progress_value = 0
    if 'stage_details' not in st.session_state:
        st.session_state.stage_details = "Esperando inicio del proceso"
    
    # Crear columnas para el layout
    col1, col2 = st.columns([2, 1])
    
    # Mostrar el estado actual del proceso en la columna derecha
    with col2:
        st.markdown("### Estado Actual")
        # Mostrar barra de progreso
        st.progress(st.session_state.progress_value)
        # Mostrar estado actual con ícono
        st.info(f"📋 {st.session_state.current_stage}")
        # Mostrar detalles adicionales
        st.caption(st.session_state.stage_details)

    with st.sidebar:
        st.header("Configuración")
        api_key = st.text_input("Ingrese su clave API de Google AI:", type="password")
        model_name = st.selectbox("Modelo", ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-1.0-pro'])
        uploaded_file = st.file_uploader("Cargar archivo de proceso", type=['txt', 'pdf', 'docx'])
        if st.button("Generar Documento"):
            if not api_key:
                st.error("Por favor, ingrese su clave API.")
                return
            if not uploaded_file:
                st.error("Por favor, suba un archivo de entrada.")
                return
            try:
                # Configurar la API
                st.session_state.current_stage = "Configurando API"
                st.session_state.progress_value = 0.1
                st.session_state.stage_details = "Inicializando conexión con la API de Google..."
                configurar_api(api_key)

                # Crear el modelo
                st.session_state.current_stage = "Inicializando modelo"
                st.session_state.progress_value = 0.2
                st.session_state.stage_details = "Cargando modelo de IA..."
                modelo = crear_modelo(model_name)

                # Cargar el archivo de entrada
                st.session_state.current_stage = "Cargando archivo"
                st.session_state.progress_value = 0.3
                st.session_state.stage_details = f"Procesando archivo: {uploaded_file.name}"
                input_text = cargar_archivo(uploaded_file)

                if input_text is None:
                    raise Exception("No se pudo cargar el archivo. Asegúrate de que el archivo es .txt, .pdf o .docx.")

                # Usar el índice proporcionado
                indice = """
1. Portada del Documento:
   Título Principal: “Documentación de Procesos ISO 9001: Guía Integral de Implementación y Control de Calidad”
   Subtítulo: “Estandarización de Procesos y Control de Calidad para Nuevas Empresas”
   Autor: Ingeniero Industrial Senior (Nombre Completo y Credenciales)
   Fecha de Elaboración: [Fecha]
   Versión del Documento: 1.0
2. Índice de Contenidos:
   Estructura el índice con numeración jerárquica detallada y descripciones concisas.
   Refleja todos los niveles de detalle con numeración clara para facilitar la navegación (1.1, 1.1.1, etc.).
3. Introducción:
   3.1. Objetivo General del Documento:
   Explica la finalidad del documento y su papel en la implementación de ISO 9001 dentro de la organización.
   3.2. Alcance y Aplicación:
   Define las áreas, departamentos y roles cubiertos por los procesos documentados.
   3.3. Metodología de Documentación:
   Describe brevemente la metodología utilizada para estructurar los procesos y asegurar la calidad y claridad de la información.
4. Definición de Responsables:
   4.1. Roles y Responsabilidades:
   Usa una tabla para detallar cada rol, sus competencias y responsabilidades dentro de cada proceso.
      | Proceso   | Responsable Principal | Rol de Apoyo  | Certificaciones Requeridas |
      |-----------|-----------------------|---------------|----------------------------|
      | Planificación | Gerente de Calidad | Ingeniero de Proceso | ISO 9001, Auditor Interno |
      | Ejecución | Supervisor de Planta | Operarios | Certificación de Maquinaria |
   4.2. Diagrama de Estructura Organizativa (Opcional):
   Incluir un diagrama que muestre las relaciones jerárquicas y de reporte dentro de cada proceso documentado (si aplica).
5. Análisis de Riesgos:
   5.1. Identificación de Riesgos:
   Realiza un análisis de riesgos por actividad, utilizando una matriz para clasificar cada riesgo por probabilidad e impacto.
   5.2. Matriz de Análisis de Riesgos:
      | Actividad del Proceso | Posible Riesgo | Probabilidad | Impacto | Medida de Mitigación |
      |-----------------------|----------------|--------------|---------|----------------------|
      | Recepción de Material | Error en especificaciones | Alta | Alto | Inspección con checklist inicial |
6. Descripción Detallada de Procesos:
   Documenta cada proceso usando numeración jerárquica para reflejar todos los pasos y subpasos.
   6.1. Proceso Principal:
         Paso 1: Describir el primer paso del subproceso con claridad.
         Paso 2: Continuar con el detalle del siguiente paso.
         Paso 3: Describir cada subpaso de manera lógica y secuencial.
7. Herramientas Utilizadas:
   7.1. Descripción de Herramientas y Recursos:
   7.2. Manuales de Uso y Mantenimiento (Si aplica):
   Incluir las guías de operación y mantenimiento de las herramientas más críticas.
8. Indicadores de Rendimiento (KPI):
   8.1. Identificación de Indicadores Clave por Proceso:
   Define los KPI específicos para cada proceso y proporciona la fórmula para su cálculo.
   8.2. Tabla de Indicadores Clave:
      | Proceso   | Indicador Clave | Fórmula de Cálculo | Valor Objetivo | Frecuencia de Medición |
      |-----------|-----------------|--------------------|----------------|------------------------|
      | Recepción de Material | Tasa de Aceptación | (# de Materiales Aceptados / Total) * 100 | 95% | Mensual |
9. Conclusión y Recomendaciones:
   Proveer un resumen conciso de los puntos críticos de cada sección.
   Incluir recomendaciones para la mejora continua y el seguimiento de los indicadores establecidos.
"""

                # Generar el documento completo utilizando el índice proporcionado
                st.session_state.current_stage = "Generando documento"
                st.session_state.progress_value = 0.5
                st.session_state.stage_details = "Analizando contenido y generando estructura del documento..."
                with st.spinner("Generando el documento..."):
                    documento_final = generar_documento(modelo, input_text, indice, model_version=model_name)
                    st.session_state.progress_value = 0.8
                    st.session_state.stage_details = "Aplicando formato y estructura ISO 9001..."

                # Guardar el documento
                st.session_state.current_stage = "Guardando documento"
                st.session_state.progress_value = 0.9
                st.session_state.stage_details = "Guardando documento final..."
                guardar_documento_word(documento_final, 'documento_procesos_ISO9001.docx')
                
                # Actualizar estado final
                st.session_state.current_stage = "¡Documento generado con éxito!"
                st.session_state.progress_value = 1.0
                st.session_state.stage_details = "El documento ha sido generado y guardado como 'documento_procesos_ISO9001.docx'"
                st.success("Documento generado exitosamente")

            except Exception as e:
                st.session_state.current_stage = "Error en el proceso"
                st.session_state.progress_value = 0
                st.session_state.stage_details = f"Error: {str(e)}"
                logging.error(f"Error en el proceso: {str(e)}")
                st.error(f"Error: {str(e)}")

    # Mostrar contenido del archivo de entrada en una pestaña
    if uploaded_file is not None:
        with col1:
            st.subheader("Contenido del archivo de entrada")
            input_text = cargar_archivo(uploaded_file)
            if input_text:
                st.text_area("", value=input_text, height=300)

if __name__ == "__main__":
    main()
